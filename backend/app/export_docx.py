# app/export_docx.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from datetime import datetime
import base64
from PIL import Image as PILImage

def _maybe_add_logo(doc: Document, logo_b64: str):
    if not logo_b64:
        return
    try:
        data = base64.b64decode(logo_b64)
        buf = BytesIO(data)
        # docx requires a file-like object
        doc.add_picture(buf, width=Inches(1.5))
    except Exception as e:
        print("DOCX logo failed:", e)

def export_docx_structured(content: str, references: list = None, file_path: str = None, filename: str = "doc", title: str = None, author: str = None, logo_b64: str = None):
    doc = Document()
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)

    if logo_b64:
        _maybe_add_logo(doc, logo_b64)

    if title:
        h = doc.add_heading(title, level=1)
        if author:
            pmeta = doc.add_paragraph(f"Author: {author}")
            pmeta.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("")

    # Body paragraphs
    for p in content.split("\n\n"):
        p = p.strip()
        if not p:
            continue
        doc.add_paragraph(p)

    # Citations
    if references:
        doc.add_page_break()
        doc.add_heading("Citations", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "#"
        hdr_cells[1].text = "Source"
        hdr_cells[2].text = "Chunk ID"
        hdr_cells[3].text = "Citation"
        for i, r in enumerate(references, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = r.get("source", "")
            row_cells[2].text = str(r.get("chunk_id", ""))
            row_cells[3].text = r.get("citation", "")

    bio = BytesIO()
    if file_path:
        doc.save(file_path)
        return str(file_path)
    else:
        doc.save(bio)
        bio.seek(0)
        return bio.read()

def export_docx(content: str, file_path: str):
    """Backward compatibility wrapper for simple exports"""
    return export_docx_structured(content=content, file_path=file_path, filename="export")
